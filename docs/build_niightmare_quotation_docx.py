from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/niightmare-quotation-lao.docx")

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(102, 102, 102)
BLACK = RGBColor(0, 0, 0)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lao UI")
    run._element.rPr.rFonts.set(qn("w:cs"), "Lao UI")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold=False, color=None, size=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return paragraph


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, 160, 180, 160, 180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=11, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_table(doc: Document, rows: list[list[str]], widths: list[float], header=True, header_fill=LIGHT):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    for row_i, row in enumerate(rows):
        tr_pr = table.rows[row_i]._tr.get_or_add_trPr()
        if row_i == 0 and header:
            hdr = OxmlElement("w:tblHeader")
            tr_pr.append(hdr)
        for col_i, text in enumerate(row):
            cell = table.cell(row_i, col_i)
            set_cell_width(cell, widths[col_i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_i == 0 and header:
                set_cell_shading(cell, header_fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            parts = str(text).split("\n")
            for part_i, part in enumerate(parts):
                if part_i:
                    p = cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.1
                run = p.add_run(part)
                set_run_font(run, size=10.5, bold=(row_i == 0 and header), color=DARK if (row_i == 0 and header) else BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ["Normal", "Body Text", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Lao UI")
        style._element.rPr.rFonts.set(qn("w:cs"), "Lao UI")
        style.font.size = Pt(11)
        style.font.color.rgb = BLACK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("NIIGHTMARE Esport - ໃບສະເໜີລາຄາໂຄງການເວັບໄຊ")
    set_run_font(run, size=9, color=MUTED)


def build() -> None:
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("ໃບສະເໜີລາຄາ")
    set_run_font(run, size=26, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("ໂຄງການເວັບໄຊ ແລະ ລະບົບຫຼັງບ້ານ NIIGHTMARE Esport")
    set_run_font(run, size=14, color=MUTED)

    add_table(
        doc,
        [
            ["ລາຍການ", "ຂໍ້ມູນ"],
            ["ຜູ້ຮັບໃບສະເໜີລາຄາ", "NIIGHTMARE Esport"],
            ["ປະເພດວຽກ", "ຂາຍງານເວັບໄຊທັງລະບົບແບບລາຄາດຽວຈົບ"],
            ["ຮູບແບບການສົ່ງມອບ", "ເວັບຫຼັກ + Admin + ລະບົບສັ່ງຊື້ເສື້ອ + ຖານຂໍ້ມູນ + Deploy"],
            ["ການດູແລຫຼັງສົ່ງມອບ", "ຟຣີ 2 ເດືອນ"],
            ["ວັນທີ", date(2026, 7, 2).strftime("%d/%m/%Y")],
            ["ສະຖານະເອກະສານ", "ສະບັບຮ່າງເພື່ອກວດລາຄາ ແລະ ປັບຕົວເລກສຸດທ້າຍ"],
        ],
        [1.8, 4.7],
    )
    add_callout(
        doc,
        "ສະຫຼຸບສັ້ນ",
        "ໂຄງການນີ້ບໍ່ແມ່ນເວັບແນະນຳທີມທຳມະດາ. ຂອບເຂດງານລວມເຖິງເວັບຫຼັກຫຼາຍໜ້າ, ລະບົບ Admin ແກ້ຂໍ້ມູນໄດ້, ລະບົບສັ່ງຊື້ເສື້ອ, ອັບໂຫຼດສະລິບ, ຈັດການອໍເດີ, ຖານຂໍ້ມູນ, ແລະການນຳຂຶ້ນໃຊ້ງານຈິງ.",
    )

    add_heading(doc, "1. ພາບລວມຂອງໂຄງການ")
    add_paragraph(
        doc,
        "ເວັບໄຊ NIIGHTMARE Esport ຖືກອອກແບບເພື່ອເປັນໜ້າບ້ານອອນລາຍຂອງທີມ esports ໃນ Lao PDR, ຮອງຮັບ Mobile Legends: Bang Bang (MLBB) ແລະ eFootball. ລະບົບຖືກພັດທະນາດ້ວຍ Next.js, TypeScript, Tailwind CSS, Supabase ແລະ Vercel ເພື່ອໃຫ້ໃຊ້ງານໄດ້ຈິງທັງໃນມືຖື ແລະ desktop.",
    )
    for item in [
        "ຮອງຮັບ 2 ພາສາ: ພາສາລາວ ແລະ English.",
        "ອອກແບບແນວ esports ແບບ dark premium ໃຫ້ກົງກັບ brand NIIGHTMARE.",
        "ເນື້ອຫາຫຼັກສາມາດແກ້ໄດ້ຜ່ານ Admin ໂດຍບໍ່ຕ້ອງ redeploy ທຸກຄັ້ງ.",
        "ເປັນລະບົບ production ທີ່ມີ domain .com ແລະໂຄງສ້າງ deploy ພ້ອມໃຊ້ງານ.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. ຂອບເຂດວຽກ: ເວັບຫຼັກ")
    sections = [
        ("2.1 ໜ້າຫຼັກ ແລະ brand experience", [
            "Hero section ສຳລັບສື່ສານຕົວຕົນຂອງທີມ.",
            "ສ່ວນ Upcoming Match ສະແດງແມັດຕໍ່ໄປແບບ broadcast card.",
            "Recent Results ແລະຂໍ້ມູນຜົນການແຂ່ງຂັນ.",
            "About Us ແບບ manifesto ຂອງທີມ ພ້ອມຂໍ້ຄວາມທີ່ Admin ແກ້ໄດ້.",
            "Partner/Sponsor strip ເພື່ອສະແດງພັນທະມິດ ແລະ sponsor.",
        ]),
        ("2.2 ໜ້າ Roster / Player Profile", [
            "ສະແດງ player cards ແລະ staff cards.",
            "ມີ profile modal ແລະໜ້າ player detail.",
            "ຮອງຮັບຂໍ້ມູນ player ເຊັ່ນ role, IGN, bio, social, achievement, Liquipedia link.",
            "ຮອງຮັບ Behind-the-Team / roster data ຜ່ານ Supabase.",
        ]),
        ("2.3 ໜ້າ Matches", [
            "ມີ scoreboard ສະຫຼຸບ W / L / win rate.",
            "ມີ filter ຕາມ game, year, result, sort ແລະ tournament.",
            "ຈັດກຸ່ມ tournament ແລະ tier color ເພື່ອໃຫ້ອ່ານງ່າຍ.",
            "ມີ Niightmare Roadmap modal ອະທິບາຍເສັ້ນທາງ H1/H2 ແລະສະຖານະປັດຈຸບັນ.",
        ]),
        ("2.4 ໜ້າ Achievements, Sponsors, Contact, Legal", [
            "Achievements ສະແດງຜົນງານ ແລະການແຂ່ງຂັນສຳຄັນ.",
            "Sponsors ສຳລັບສະແດງ sponsor ແລະການຕິດຕໍ່ຮ່ວມງານ.",
            "Contact ແລະ social links ສຳລັບ fan, sponsor, ແລະຜູ້ຕິດຕໍ່.",
            "Privacy Policy ແລະ Terms ເພື່ອໃຫ້ເວັບດູເປັນທາງການ.",
        ]),
        ("2.5 SEO, responsive, PWA", [
            "SEO metadata, Open Graph image, Twitter image, sitemap, robots.",
            "Responsive layout ສຳລັບ mobile, tablet, desktop.",
            "Install prompt / PWA manifest ເພື່ອໃຫ້ fan ຕິດຕັ້ງເວັບເປັນ app ໄດ້.",
            "Self-hosted fonts ເພື່ອຫຼີກລ້ຽງ build hang ແລະຮອງຮັບ Lao font.",
        ]),
    ]
    for heading, bullets in sections:
        add_heading(doc, heading, 2)
        for item in bullets:
            add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "3. ຂອບເຂດວຽກ: ລະບົບ Admin / ຫຼັງບ້ານ")
    add_paragraph(doc, "ລະບົບ Admin ຖືກສ້າງຂຶ້ນເພື່ອໃຫ້ທີມສາມາດຈັດການເນື້ອຫາ ແລະອໍເດີໄດ້ດ້ວຍຕົນເອງ ໂດຍບໍ່ຈຳເປັນຕ້ອງແກ້ code ທຸກຄັ້ງ.")
    for heading, bullets in [
        ("3.1 ຄວາມປອດໄພ ແລະ session", [
            "Admin login ສຳລັບກັ່ນກອງຜູ້ເຂົ້າໃຊ້.",
            "Session cookie ສຳລັບຈື່ຈຳການ login ໃນອຸປະກອນທີ່ໃຊ້ງານ.",
            "ຮອງຮັບ 2FA/TOTP ຖ້າຕັ້ງຄ່າ secret ໃນ environment.",
            "ມີ API protection ສຳລັບ route ທີ່ແກ້ຂໍ້ມູນ.",
        ]),
        ("3.2 ຈັດການເນື້ອຫາເວັບ", [
            "Home/About Us editor.",
            "Roadmap editor ສຳລັບໜ້າ Matches.",
            "Roster editor ແກ້ຂໍ້ມູນ player/staff.",
            "Matches editor ແກ້ຜົນແຂ່ງຂັນ ແລະ upcoming match.",
            "Achievements editor.",
            "Sponsors editor.",
            "Shop editor ແກ້ລາຄາ, size, stock, bank/QR, courier, contact link.",
            "ລະບົບ upload image ສຳລັບ media ແລະຮູບປະກອບ.",
        ]),
        ("3.3 Community moderation", [
            "ຮອງຮັບ fan auth ຜ່ານ Google / magic link.",
            "ມີ comment route ສຳລັບ player ແລະ team support.",
            "Admin ສາມາດກວດ ແລະ moderate comments ໄດ້.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for item in bullets:
            add_bullet(doc, item)

    add_heading(doc, "4. ຂອບເຂດວຽກ: ລະບົບ Shop / Jersey Ordering")
    add_paragraph(doc, "ລະບົບ Shop ເປັນລະບົບສັ່ງຊື້ເສື້ອແບບ on-site ໂດຍບໍ່ບັງຄັບ login. ຜູ້ຊື້ສາມາດເລືອກ size, ກອກຂໍ້ມູນຈັດສົ່ງ, ຈ່າຍເງິນຜ່ານ QR/ໂອນເງິນ, ແລະອັບໂຫຼດສະລິບເພື່ອໃຫ້ທີມກວດສອບ.")
    for heading, bullets in [
        ("4.1 ຝັ່ງຜູ້ຊື້", [
            "ເລືອກ size S-4XL ແລະຈຳນວນຫຼາຍ size ໃນອໍເດີດຽວ.",
            "ຄຳນວນລາຄາ server-side ເພື່ອປ້ອງກັນການປ່ຽນລາຄາຈາກ browser.",
            "ກອກຊື່, ເບີໂທ/WhatsApp, courier, province/city/branch.",
            "ສ້າງ reference code ເຊັ່ນ NM-XXXXX ໃຫ້ຜູ້ຊື້ໃສ່ໃນ transfer note.",
            "Payment popup ພ້ອມ QR, ຍອດໂອນ, reference code, ແລະ upload slip.",
            "My Orders ບັນທຶກອໍເດີໃນ localStorage ແລະ sync status ຈາກ server.",
            "ມີ 24-hour pay window ສຳລັບອໍເດີທີ່ຍັງບໍ່ໄດ້ຈ່າຍ.",
            "ສະແດງຮູບຈັດສົ່ງເມື່ອ Admin ອັບໂຫຼດໃຫ້.",
        ]),
        ("4.2 ຝັ່ງ Admin Orders", [
            "ເບິ່ງອໍເດີຕາມ status: ລໍຖ້າຊຳລະ, ກຳລັງກວດ, ຈ່າຍແລ້ວ/ຍົກເລີກ, ສົ່ງແລ້ວ.",
            "ສະແດງ ref code, ຍອດເງິນ, ຮູບສະລິບ ເພື່ອກວດເງິນໄດ້ໄວ.",
            "ຄົ້ນຫາດ້ວຍ ref code, ຊື່, email, ເບີໂທ.",
            "Sort ໃໝ່ຫາເກົ່າ / ເກົ່າຫາໃໝ່.",
            "Duplicate flag ເມື່ອມີເບີໂທ ຫຼື email ຊ້ຳ.",
            "Quick advance status: paid -> verified -> shipped.",
            "Upload shipping image ໃຫ້ຜູ້ຊື້ເຫັນ.",
            "Sales report ສະຫຼຸບຍອດຂາຍ, ຈຳນວນເສື້ອ, ແລະຍອດຕາມ size.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for item in bullets:
            add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "5. ໂຄງສ້າງ Backend, Database, Deploy")
    add_table(
        doc,
        [
            ["ສ່ວນລະບົບ", "ລາຍລະອຽດ"],
            ["Frontend", "Next.js 15 App Router, React, TypeScript, Tailwind CSS, Framer Motion, responsive UI."],
            ["Database", "Supabase tables ສຳລັບ players, matches, upcoming match, site settings, shop orders, push subscriptions."],
            ["Storage", "Supabase Storage uploads bucket ສຳລັບ image uploads, payment slips, shipping images."],
            ["API", "Next.js API routes ສຳລັບ admin data, upload, orders, comments, auth, push notification."],
            ["Deploy", "Vercel production deploy, custom domain niightmareesport.com, Cloudflare DNS."],
            ["Notification", "Web Push ສຳລັບແຈ້ງ Admin ເມື່ອຜູ້ຊື້ແຈ້ງໂອນເງິນ."],
            ["Email", "Formspree ສຳລັບ email notification ຕາມ flow ທີ່ຕັ້ງຄ່າ."],
            ["Security", "Admin auth, protected API routes, validation, production-oriented configuration."],
        ],
        [1.5, 5.0],
    )

    add_heading(doc, "6. ສິ່ງທີ່ສົ່ງມອບ")
    for item in [
        "Source code ຂອງເວັບທັງໝົດ.",
        "ເວັບ production ທີ່ເປີດໃຊ້ງານຜ່ານ niightmareesport.com.",
        "Admin panel ສຳລັບຈັດການ content, media, shop config, ແລະ orders.",
        "Database schema ແລະ integration ກັບ Supabase.",
        "Upload flow ສຳລັບ image, payment slip, shipping image.",
        "PWA/manifest/icons ສຳລັບ public site ແລະ admin.",
        "ການຕັ້ງຄ່າ deploy ແລະ production environment ທີ່ຈຳເປັນ.",
        "ການອະທິບາຍວິທີໃຊ້ Admin ແລະດູແລເບື້ອງຕົ້ນ.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. ລາຄາປະເມີນຕາມຂອບເຂດງານ")
    add_paragraph(doc, "ຕາຕະລາງນີ້ເປັນລາຄາປະເມີນຕາມມູນຄ່າວຽກປົກກະຕິ. ຜູ້ພັດທະນາສາມາດປັບເປັນລາຄາພິເສດໃຫ້ NIIGHTMARE Esport ໄດ້ຕາມຄວາມເໝາະສົມ.")
    add_table(
        doc,
        [
            ["ຫມວດວຽກ", "ລາຍລະອຽດ", "ມູນຄ່າປະເມີນ"],
            ["ເວັບຫຼັກ / UX UI", "ອອກແບບໜ້າຫຼັກ, roster, matches, achievements, sponsors, contact, responsive, bilingual", "28,000,000 LAK"],
            ["Admin / CMS", "Login, content editors, media upload, roster/match/sponsor/shop management", "30,000,000 LAK"],
            ["Shop / Order system", "ສັ່ງເສື້ອ, ຄຳນວນລາຄາ, QR/slip, My Orders, order status, shipping image", "30,000,000 LAK"],
            ["Backend / Database / Security", "Supabase, storage, API routes, validation, auth, push notification, deploy setup", "22,000,000 LAK"],
            ["QA / ສອນໃຊ້ / Support", "ກວດກາ production, ແກ້ bug ເບື້ອງຕົ້ນ, ສອນໃຊ້ Admin, ດູແລຟຣີ 2 ເດືອນ", "10,000,000 LAK"],
            ["ລວມມູນຄ່າປົກກະຕິ", "", "120,000,000 LAK"],
            ["ສ່ວນຫຼຸດພິເສດຄົນສະນິດ", "ລໍຖ້າກຳນົດ", "................ LAK"],
            ["ລາຄາສະເໜີສຸດທ້າຍ", "ລາຄາດຽວຈົບ", "................ LAK"],
        ],
        [1.6, 3.5, 1.4],
        header_fill=PALE_BLUE,
    )
    add_callout(doc, "ຂໍ້ແນະນຳເລື່ອງລາຄາ", "ຖ້າເປັນລາຄາພິເສດສຳລັບຄົນສະນິດ, ສາມາດປັບຈາກມູນຄ່າປົກກະຕິລົງມາໄດ້. ແຕ່ບໍ່ຄວນຕັດຈົນຕ່ຳເກີນໄປ ເພາະຂອບເຂດວຽກມີທັງ frontend, backend, admin, database, deploy ແລະການດູແລຫຼັງສົ່ງມອບ.")

    add_heading(doc, "8. ຄ່າໃຊ້ຈ່າຍພາຍນອກ")
    add_paragraph(doc, "ຄ່າໃຊ້ຈ່າຍພາຍນອກເປັນຄ່າ platform/subscription ທີ່ບໍ່ລວມໃນຄ່າພັດທະນາເວັບ. ຜູ້ຈ້າງເປັນຜູ້ຮັບຜິດຊອບໂດຍກົງ. ການຄຳນວນ LAK ດ້ານລຸ່ມໃຊ້ອັດຕາປະເມີນ 1 USD = 22,500 LAK ເພື່ອໃຫ້ເຫັນພາບລວມ.")
    add_table(
        doc,
        [
            ["ລາຍການ", "ຄ່າໃຊ້ຈ່າຍ", "ປະມານເປັນ LAK", "ໝາຍເຫດ"],
            ["Domain .com", "11 USD", "247,500 LAK", "ປົກກະຕິເປັນລາຍປີ"],
            ["Claude Pro", "20 USD", "450,000 LAK", "ລາຍເດືອນ ຖ້າຍັງຈຳເປັນໃຊ້"],
            ["Codex Pro", "20 USD", "450,000 LAK", "ລາຍເດືອນ ຖ້າຍັງຈຳເປັນໃຊ້"],
            ["Vercel / Supabase / Formspree", "ຂຶ້ນກັບ usage", "ລໍຖ້າໃບບິນຈິງ", "ອາດຟຣີໃນຊ່ວງເລີ່ມຕົ້ນ ຫຼືມີຄ່າໃຊ້ຈ່າຍເມື່ອໃຊ້ຫຼາຍ"],
        ],
        [1.7, 1.2, 1.4, 2.2],
    )

    add_heading(doc, "9. ເງື່ອນໄຂການດູແລຟຣີ 2 ເດືອນ")
    for heading, bullets in [
        ("ລວມໃນການດູແລຟຣີ", [
            "ແກ້ bug ທີ່ເກີດຈາກລະບົບທີ່ສົ່ງມອບ.",
            "ຊ່ວຍ deploy ແລະກວດ production ເມື່ອມີການແກ້ໄຂຈຳເປັນ.",
            "ຊ່ວຍຕັ້ງຄ່າ Admin, Shop, QR, bank info, courier, ແລະ content ເບື້ອງຕົ້ນ.",
            "ສອນການໃຊ້ Admin ແລະວິທີກວດອໍເດີ.",
        ]),
        ("ບໍ່ລວມໃນການດູແລຟຣີ", [
            "ຟີເຈີໃໝ່ທີ່ບໍ່ຢູ່ໃນ scope ເດີມ.",
            "ການ redesign ໃຫຍ່ ຫຼືປ່ຽນ brand direction.",
            "ລະບົບ payment gateway ອັດຕະໂນມັດ.",
            "3D model ແບບ custom .glb ຫຼືງານ animation ພິເສດ.",
            "ຄ່າ subscription / platform / domain / storage ຂອງບໍລິການພາຍນອກ.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for item in bullets:
            add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "10. ຂໍ້ຍົກເວັ້ນ ແລະ ຂໍ້ຈຳກັດ")
    for item in [
        "ລາຄານີ້ບໍ່ລວມຄ່າ platform ພາຍນອກ ເຊັ່ນ domain, AI subscription, Vercel, Supabase, Formspree ຫຼືຄ່າບໍລິການອື່ນ.",
        "ລະບົບ payment ໃນປັດຈຸບັນແມ່ນ manual slip verification, ບໍ່ແມ່ນ gateway ທີ່ກວດເງິນອັດຕະໂນມັດ.",
        "ຖ້າຕ້ອງການ feature ໃໝ່ຫຼັງຈາກສົ່ງມອບ ຄວນປະເມີນເປັນວຽກເພີ່ມ.",
        "ຖ້າບໍລິການພາຍນອກມີບັນຫາ ເຊັ່ນ Supabase, Vercel, Cloudflare ລົ້ມ ຫຼືປ່ຽນ policy, ການແກ້ໄຂອາດຕ້ອງປະເມີນເພີ່ມຕາມຄວາມຍາກ.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. ການຢືນຢັນ ແລະ ລາຍເຊັນ")
    add_paragraph(doc, "ຖ້າຜູ້ຈ້າງເຫັນດີກັບ scope ແລະລາຄາສຸດທ້າຍ, ທັງສອງຝ່າຍສາມາດຢືນຢັນດ້ວຍການລົງຊື່ ຫຼືການຕົກລົງເປັນລາຍລັກອັກສອນ.")
    add_table(
        doc,
        [
            ["ຝ່າຍຜູ້ຈ້າງ", "ຝ່າຍຜູ້ພັດທະນາ"],
            ["ຊື່: ______________________________\nລາຍເຊັນ: _________________________\nວັນທີ: ____ / ____ / ______", "ຊື່: ______________________________\nລາຍເຊັນ: _________________________\nວັນທີ: ____ / ____ / ______"],
        ],
        [3.25, 3.25],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
